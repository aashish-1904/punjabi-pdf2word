import streamlit as st
import os
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
import io
import requests
from urllib.parse import urlparse

# Set page config
st.set_page_config(
    page_title="Punjabi Text Converter",
    page_icon="📄",
    layout="centered"
)

# Custom CSS
st.markdown("""
    <style>
    .stButton>button {
        background: linear-gradient(90deg, #009688 60%, #4dd0e1 100%);
        color: white;
        border: none;
        padding: 0.5rem 1rem;
        border-radius: 0.5rem;
        font-weight: 600;
    }
    .stButton>button:hover {
        background: linear-gradient(90deg, #00796b 60%, #26c6da 100%);
    }
    .main {
        background: linear-gradient(120deg, #f5f5f5 60%, #e0f7fa 100%);
    }
    </style>
    """, unsafe_allow_html=True)

# Global font status
FONTS_REGISTERED = {
    'gurmukhi_regular': False,
    'gurmukhi_bold': False,
    'helvetica_available': True
}

@st.cache_data
def download_font(font_url, font_name):
    """Download and cache a font file"""
    try:
        response = requests.get(font_url, timeout=30)
        response.raise_for_status()
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.ttf') as font_file:
            font_file.write(response.content)
            return font_file.name
    except Exception as e:
        st.warning(f"Could not download {font_name}: {e}")
        return None

def register_fonts():
    """Register all required fonts"""
    global FONTS_REGISTERED
    
    try:
        # Download and register Gurmukhi Regular
        regular_url = "https://github.com/notofonts/notofonts.github.io/raw/main/fonts/NotoSansGurmukhi/hinted/ttf/NotoSansGurmukhi-Regular.ttf"
        regular_path = download_font(regular_url, "Noto Sans Gurmukhi Regular")
        
        if regular_path:
            pdfmetrics.registerFont(TTFont('NotoSansGurmukhi', regular_path))
            FONTS_REGISTERED['gurmukhi_regular'] = True
            
        # Download and register Gurmukhi Bold
        bold_url = "https://github.com/notofonts/notofonts.github.io/raw/main/fonts/NotoSansGurmukhi/hinted/ttf/NotoSansGurmukhi-Bold.ttf"
        bold_path = download_font(bold_url, "Noto Sans Gurmukhi Bold")
        
        if bold_path:
            pdfmetrics.registerFont(TTFont('NotoSansGurmukhi-Bold', bold_path))
            FONTS_REGISTERED['gurmukhi_bold'] = True
            
    except Exception as e:
        st.error(f"Font registration error: {e}")

def is_gurmukhi_text(text):
    """Check if text contains Gurmukhi characters"""
    if not text:
        return False
    gurmukhi_range = range(0x0A00, 0x0A7F)
    return any(ord(char) in gurmukhi_range for char in text)

def get_text_formatting(run):
    """Extract formatting from a Word run"""
    formatting = {
        'bold': False,
        'italic': False,
        'underline': False,
        'size': 12,
        'color': '#000000'
    }
    
    try:
        # Check bold
        if (hasattr(run, 'bold') and run.bold is True) or \
           (hasattr(run, 'font') and hasattr(run.font, 'bold') and run.font.bold is True):
            formatting['bold'] = True
            
        # Check italic
        if (hasattr(run, 'italic') and run.italic is True) or \
           (hasattr(run, 'font') and hasattr(run.font, 'italic') and run.font.italic is True):
            formatting['italic'] = True
            
        # Check underline
        if (hasattr(run, 'underline') and run.underline is True) or \
           (hasattr(run, 'font') and hasattr(run.font, 'underline') and run.font.underline is True):
            formatting['underline'] = True
            
        # Get font size
        if hasattr(run, 'font') and run.font.size and hasattr(run.font.size, 'pt'):
            formatting['size'] = run.font.size.pt
            
        # Get color
        if hasattr(run, 'font') and run.font.color and hasattr(run.font.color, 'rgb') and run.font.color.rgb:
            rgb_val = run.font.color.rgb
            formatting['color'] = f"#{rgb_val:06x}"
            
    except Exception:
        pass
        
    return formatting

def get_best_font(text, is_bold=False, is_italic=False):
    """Get the best available font for the given text and formatting"""
    if is_gurmukhi_text(text):
        # Punjabi text - for italic, always use Helvetica for proper rendering
        if is_italic:
            # Always use Helvetica variants for italic Punjabi text
            if is_bold:
                return 'Helvetica-BoldOblique'
            else:
                return 'Helvetica-Oblique'
        elif is_bold and FONTS_REGISTERED['gurmukhi_bold']:
            return 'NotoSansGurmukhi-Bold'
        elif FONTS_REGISTERED['gurmukhi_regular']:
            return 'NotoSansGurmukhi'
        else:
            # Fallback to Helvetica
            if is_bold:
                return 'Helvetica-Bold'
            else:
                return 'Helvetica'
    else:
        # Non-Punjabi text - use Helvetica variants
        if is_bold and is_italic:
            return 'Helvetica-BoldOblique'
        elif is_bold:
            return 'Helvetica-Bold'
        elif is_italic:
            return 'Helvetica-Oblique'
        else:
            return 'Helvetica'

def get_paragraph_alignment(paragraph):
    """Get ReportLab alignment from Word paragraph"""
    alignment_map = {
        0: TA_LEFT,      # WD_ALIGN_PARAGRAPH.LEFT
        1: TA_CENTER,    # WD_ALIGN_PARAGRAPH.CENTER
        2: TA_RIGHT,     # WD_ALIGN_PARAGRAPH.RIGHT
        3: TA_JUSTIFY,   # WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    if hasattr(paragraph, 'alignment') and paragraph.alignment is not None:
        return alignment_map.get(paragraph.alignment, TA_LEFT)
    return TA_LEFT

def hex_to_reportlab_color(hex_color):
    """Convert hex color to ReportLab color"""
    try:
        if hex_color.startswith('#'):
            hex_color = hex_color[1:]
        r = int(hex_color[0:2], 16) / 255.0
        g = int(hex_color[2:4], 16) / 255.0
        b = int(hex_color[4:6], 16) / 255.0
        return colors.Color(r, g, b)
    except:
        return colors.black

def convert_docx_to_pdf(docx_bytes, output_filename):
    """Convert DOCX to PDF with comprehensive formatting preservation"""
    try:
        # Register fonts
        register_fonts()
        
        # Create temporary DOCX file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_docx:
            temp_docx.write(docx_bytes)
            temp_docx_path = temp_docx.name

        # Read the document
        doc = Document(temp_docx_path)
        
        # Analyze document formatting
        stats = {
            'bold_runs': 0,
            'italic_runs': 0,
            'underline_runs': 0,
            'colored_runs': 0,
            'gurmukhi_runs': 0,
            'font_sizes': set()
        }
        
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.text.strip():
                    fmt = get_text_formatting(run)
                    if fmt['bold']:
                        stats['bold_runs'] += 1
                    if fmt['italic']:
                        stats['italic_runs'] += 1
                    if fmt['underline']:
                        stats['underline_runs'] += 1
                    if fmt['color'] != '#000000':
                        stats['colored_runs'] += 1
                    if is_gurmukhi_text(run.text):
                        stats['gurmukhi_runs'] += 1
                    stats['font_sizes'].add(fmt['size'])
        
        # Display analysis
        st.success(f"""
        📊 **Document Analysis:**
        - Bold text runs: {stats['bold_runs']}
        - Italic text runs: {stats['italic_runs']}
        - Underlined text runs: {stats['underline_runs']}
        - Colored text runs: {stats['colored_runs']}
        - Punjabi text runs: {stats['gurmukhi_runs']}
        - Font sizes: {sorted(list(stats['font_sizes']))}
        
        🔤 **Font Status:**
        - Gurmukhi Regular: {'✅' if FONTS_REGISTERED['gurmukhi_regular'] else '❌'}
        - Gurmukhi Bold: {'✅' if FONTS_REGISTERED['gurmukhi_bold'] else '❌'}
        - Helvetica variants: ✅
        """)
        
        # Create PDF
        buffer = io.BytesIO()
        pdf_doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            topMargin=1*inch,
            bottomMargin=1*inch,
            leftMargin=1*inch,
            rightMargin=1*inch
        )
        
        styles = getSampleStyleSheet()
        story = []
        
        # Process each paragraph
        for para_idx, paragraph in enumerate(doc.paragraphs):
            if not paragraph.text.strip():
                continue
                
            # Get paragraph properties
            alignment = get_paragraph_alignment(paragraph)
            style_name = paragraph.style.name if paragraph.style else 'Normal'
            
            # Determine base font size
            base_size = 12
            if 'Heading 1' in style_name:
                base_size = 18
            elif 'Heading 2' in style_name:
                base_size = 16
            elif 'Heading 3' in style_name:
                base_size = 14
            elif 'Title' in style_name:
                base_size = 20
            
            # Handle list items
            is_list_item = 'List' in style_name
            list_prefix = ""
            if is_list_item:
                if 'Bullet' in style_name:
                    list_prefix = "● "  # Use filled circle bullet (U+25CF) which renders better
                elif 'Number' in style_name:
                    # Simple numbering - could be enhanced
                    list_number = sum(1 for i, p in enumerate(doc.paragraphs[:para_idx]) 
                                    if 'List Number' in (p.style.name if p.style else ''))
                    list_prefix = f"{list_number + 1}. "
            
            # Check if paragraph has uniform formatting
            runs_with_text = [run for run in paragraph.runs if run.text.strip()]
            if not runs_with_text:
                continue
                
            # Get formatting of first run
            first_fmt = get_text_formatting(runs_with_text[0])
            uniform_formatting = True
            
            for run in runs_with_text[1:]:
                run_fmt = get_text_formatting(run)
                if (run_fmt['bold'] != first_fmt['bold'] or 
                    run_fmt['italic'] != first_fmt['italic'] or
                    run_fmt['underline'] != first_fmt['underline']):
                    uniform_formatting = False
                    break
            
            if uniform_formatting:
                # Single paragraph with uniform formatting
                font_name = get_best_font(paragraph.text, first_fmt['bold'], first_fmt['italic'])
                text_color = hex_to_reportlab_color(first_fmt['color'])
                
                # Adjust spacing for lists
                space_after = 3 if is_list_item else 6
                left_indent = 20 if is_list_item else 0
                
                para_style = ParagraphStyle(
                    f'Para_{para_idx}',
                    parent=styles['Normal'],
                    fontName=font_name,
                    fontSize=max(first_fmt['size'], base_size),
                    alignment=alignment,
                    textColor=text_color,
                    spaceAfter=space_after,
                    leftIndent=left_indent,
                    leading=max(first_fmt['size'], base_size) * 1.2
                )
                
                # Handle underline and add list prefix
                text_content = list_prefix + paragraph.text
                if first_fmt['underline']:
                    text_content = f"<u>{text_content}</u>"
                
                story.append(Paragraph(text_content, para_style))
                
                # Debug info
                if first_fmt['bold'] or first_fmt['italic'] or is_list_item:
                    formatting_info = []
                    if first_fmt['bold']:
                        formatting_info.append("BOLD")
                    if first_fmt['italic']:
                        formatting_info.append("ITALIC")
                    if is_list_item:
                        formatting_info.append("LIST")
                    st.write(f"🎯 Uniform: {font_name} [{', '.join(formatting_info)}] - {text_content[:50]}...")
                    
            else:
                # Mixed formatting - handle each run separately
                if is_list_item:
                    # For list items with mixed formatting, add prefix to first run
                    first_run_processed = False
                
                for run_idx, run in enumerate(runs_with_text):
                    if not run.text.strip():
                        continue
                        
                    run_fmt = get_text_formatting(run)
                    font_name = get_best_font(run.text, run_fmt['bold'], run_fmt['italic'])
                    text_color = hex_to_reportlab_color(run_fmt['color'])
                    
                    # Add list prefix to first run only
                    run_text = run.text
                    if is_list_item and not first_run_processed:
                        run_text = list_prefix + run_text
                        first_run_processed = True
                    
                    # Adjust spacing and indentation for lists
                    space_after = 1 if is_list_item else 2
                    left_indent = 20 if is_list_item else 0
                    
                    run_style = ParagraphStyle(
                        f'Run_{para_idx}_{run_idx}',
                        parent=styles['Normal'],
                        fontName=font_name,
                        fontSize=max(run_fmt['size'], base_size),
                        alignment=alignment,
                        textColor=text_color,
                        spaceAfter=space_after,
                        leftIndent=left_indent,
                        leading=max(run_fmt['size'], base_size) * 1.2
                    )
                    
                    # Handle underline
                    text_content = run_text
                    if run_fmt['underline']:
                        text_content = f"<u>{run_text}</u>"
                    
                    story.append(Paragraph(text_content, run_style))
                    
                    # Debug info
                    if run_fmt['bold'] or run_fmt['italic'] or is_list_item:
                        formatting_info = []
                        if run_fmt['bold']:
                            formatting_info.append("BOLD")
                        if run_fmt['italic']:
                            formatting_info.append("ITALIC")
                        if is_list_item:
                            formatting_info.append("LIST")
                        st.write(f"🎯 Mixed: {font_name} [{', '.join(formatting_info)}] - {text_content[:30]}...")
                
                # Add paragraph spacing for mixed formatting
                if not is_list_item:
                    story.append(Spacer(1, 6))
        
        # Process tables with enhanced formatting
        for table in doc.tables:
            if not table.rows:
                continue
                
            table_data = []
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    # Process each cell with formatting
                    cell_content = ""
                    for para in cell.paragraphs:
                        if para.text.strip():
                            # Check if cell has formatting
                            has_formatting = False
                            for run in para.runs:
                                if run.text.strip():
                                    fmt = get_text_formatting(run)
                                    if fmt['bold'] or fmt['italic'] or fmt['underline']:
                                        has_formatting = True
                                        break
                            
                            if has_formatting:
                                # Build formatted cell content
                                para_content = ""
                                for run in para.runs:
                                    if run.text.strip():
                                        fmt = get_text_formatting(run)
                                        run_text = run.text
                                        
                                        # Apply formatting tags
                                        if fmt['underline']:
                                            run_text = f"<u>{run_text}</u>"
                                        if fmt['italic']:
                                            run_text = f"<i>{run_text}</i>"
                                        if fmt['bold']:
                                            run_text = f"<b>{run_text}</b>"
                                        
                                        para_content += run_text
                                cell_content += para_content + " "
                            else:
                                # Plain text
                                cell_content += para.text + " "
                    
                    row_data.append(cell_content.strip())
                table_data.append(row_data)
            
            if table_data:
                # Create table with proper styling
                pdf_table = Table(table_data)
                
                # Enhanced table styling
                table_style = [
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('FONTNAME', (0, 0), (-1, -1), 'NotoSansGurmukhi' if FONTS_REGISTERED['gurmukhi_regular'] else 'Helvetica'),
                    ('FONTSIZE', (0, 0), (-1, -1), 11),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                    ('LEFTPADDING', (0, 0), (-1, -1), 8),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 8),
                    ('TOPPADDING', (0, 0), (-1, -1), 6),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ]
                
                # Make header row bold and highlighted
                if len(table_data) > 1:
                    table_style.extend([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                        ('FONTNAME', (0, 0), (-1, 0), 'NotoSansGurmukhi-Bold' if FONTS_REGISTERED['gurmukhi_bold'] else 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ])
                
                pdf_table.setStyle(TableStyle(table_style))
                story.append(pdf_table)
                story.append(Spacer(1, 12))
                
                # Debug info
                st.write(f"🎯 Table: {len(table_data)} rows, {len(table_data[0]) if table_data else 0} columns processed")
        
        # Build PDF
        pdf_doc.build(story)
        
        # Cleanup
        os.unlink(temp_docx_path)
        
        buffer.seek(0)
        return buffer.getvalue()
        
    except Exception as e:
        st.error(f"Conversion error: {str(e)}")
        import traceback
        st.error(f"Details: {traceback.format_exc()}")
        return None

# Main UI
st.title("ਪੰਜਾਬੀ ਟੈਕਸਟ ਕਨਵਰਟਰ")
st.markdown("**Punjabi Text Converter** - Convert Word documents to PDF with perfect formatting")

st.info("""
**🎯 Features:**
- ✅ Authentic Punjabi (Gurmukhi) font rendering
- ✅ Bold and italic text preservation
- ✅ Color and font size preservation
- ✅ Table formatting support
- ✅ Mixed language support (Punjabi + English)
- ✅ Automatic font fallback system
""")

# File uploader
uploaded_file = st.file_uploader("Choose a Word document (.docx)", type=['docx'])

if uploaded_file is not None:
    try:
        with st.spinner('🔄 Converting document... Please wait...'):
            pdf_bytes = convert_docx_to_pdf(uploaded_file.getvalue(), uploaded_file.name)
            
            if pdf_bytes:
                st.success('✅ Conversion completed successfully!')
                
                # Download buttons
                col1, col2 = st.columns(2)
                
                with col1:
                    st.download_button(
                        label="📄 Preview PDF",
                        data=pdf_bytes,
                        file_name=f"preview_{Path(uploaded_file.name).stem}.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                
                with col2:
                    st.download_button(
                        label="⬇️ Download PDF",
                        data=pdf_bytes,
                        file_name=f"{Path(uploaded_file.name).stem}.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                
                # File info
                st.info(f"📊 Original: {uploaded_file.name} ({len(uploaded_file.getvalue())/1024:.1f} KB)")
                st.info(f"📄 PDF: {len(pdf_bytes)/1024:.1f} KB")
                
    except Exception as e:
        st.error(f'❌ Error: {str(e)}')
        st.error("Please ensure your document is a valid Word file and try again.") 
